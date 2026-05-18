from datetime import date
from pathlib import Path
from math import ceil

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "Docs"
IMG = DOCS / "img"
SCHEMATICS = ROOT / "Electrical_schematics"
OUT = DOCS / "Dokumentacja_AHRS.docx"


POLISH_REPLACEMENTS = {
    "Krakow": "Kraków",
    "tresci": "treści",
    "urzadzen": "urządzeń",
    "urzadzenia": "urządzenia",
    "urzadzenie": "urządzenie",
    "urzadzeniu": "urządzeniu",
    "urzadzeniami": "urządzeniami",
    "Glowny": "Główny",
    "glowny": "główny",
    "glowne": "główne",
    "Glowne": "Główne",
    "warstwe": "warstwę",
    "warstwa": "warstwa",
    "uzytkownika": "użytkownika",
    "uzytkownikiem": "użytkownikiem",
    "uzywany": "używany",
    "uzywana": "używana",
    "uzyte": "użyte",
    "uzyta": "użyta",
    "uzyty": "użyty",
    "uzyciem": "użyciem",
    "uzyciu": "użyciu",
    "uzywac": "używać",
    "uzywanym": "używanym",
    "uzytkowe": "użytkowe",
    "wyswietlaczu": "wyświetlaczu",
    "wyswietlacza": "wyświetlacza",
    "wyswietlacz": "wyświetlacz",
    "wyswietla": "wyświetla",
    "wyswietlania": "wyświetlania",
    "wyswietlaniem": "wyświetlaniem",
    "wyswietlane": "wyświetlane",
    "wyswietlac": "wyświetlać",
    "czujnikow": "czujników",
    "czujnikami": "czujnikami",
    "czujnik": "czujnik",
    "zyroskop": "żyroskop",
    "zyroskopu": "żyroskopu",
    "zyroskopem": "żyroskopem",
    "katow": "kątów",
    "katy": "kąty",
    "kat": "kąt",
    "kata": "kąta",
    "przechylenia": "przechylenia",
    "pochylenia": "pochylenia",
    "spojnosci": "spójności",
    "spojnosc": "spójność",
    "rownania": "równania",
    "rownan": "równań",
    "rownanie": "równanie",
    "rownolegle": "równolegle",
    "rowniez": "również",
    "dokladniej": "dokładniej",
    "dokladnosc": "dokładność",
    "dokladnego": "dokładnego",
    "dokladna": "dokładna",
    "wlasny": "własny",
    "wlasna": "własna",
    "wlasnego": "własnego",
    "autorska": "autorska",
    "mozliwosci": "możliwości",
    "mozliwosc": "możliwość",
    "moze": "może",
    "moze": "może",
    "ktory": "który",
    "ktora": "która",
    "ktore": "które",
    "ktorych": "których",
    "ktorym": "którym",
    "ktorymi": "którymi",
    "ktorego": "którego",
    "rozdzialy": "rozdziały",
    "rozdzialow": "rozdziałów",
    "czesc": "część",
    "czesci": "części",
    "czestotliwosc": "częstotliwość",
    "czestotliwosci": "częstotliwości",
    "probkowania": "próbkowania",
    "predkosci": "prędkości",
    "predkosc": "prędkość",
    "pamieci": "pamięci",
    "pamiec": "pamięć",
    "odswiezanie": "odświeżanie",
    "odswiezania": "odświeżania",
    "odswieza": "odświeża",
    "odswiezac": "odświeżać",
    "przerysowanie": "przerysowanie",
    "wspolnej": "wspólnej",
    "wspolne": "wspólne",
    "wspolna": "wspólna",
    "zrodlo": "źródło",
    "zrodlowe": "źródłowe",
    "zrodlowy": "źródłowy",
    "srodowisko": "środowisko",
    "srednia": "średnia",
    "miedzy": "między",
    "niepewnosc": "niepewność",
    "pomiary": "pomiary",
    "pomiarow": "pomiarów",
    "pomiaru": "pomiaru",
    "wartosc": "wartość",
    "wartosci": "wartości",
    "zawartosc": "zawartość",
    "zawartosci": "zawartości",
    "zawiera": "zawiera",
    "zawijany": "zawijany",
    "przyjetego": "przyjętego",
    "przyjeto": "przyjęto",
    "przeliczenie": "przeliczenie",
    "przelicznik": "przelicznik",
    "wejscia": "wejścia",
    "wyjscia": "wyjścia",
    "polaczen": "połączeń",
    "polaczenia": "połączenia",
    "podlaczyc": "podłączyć",
    "podlaczeniu": "podłączeniu",
    "wlaczania": "włączania",
    "wlaczyc": "włączyć",
    "wlaczany": "włączany",
    "wlaczona": "włączona",
    "wlaczone": "włączone",
    "wlacznik": "włącznik",
    "zasilania": "zasilania",
    "uklad": "układ",
    "ukladu": "układu",
    "uklady": "układy",
    "ukladow": "układów",
    "ukladem": "układem",
    "ksztalt": "kształt",
    "zlozyc": "złożyć",
    "pelni": "pełni",
    "pelna": "pełna",
    "pelny": "pełny",
    "rown": "równ",
    "Rownania": "Równania",
}

POST_POLISH_REPLACEMENTS = {
    "kątalog": "katalog",
    "przykladow": "przykładow",
    "Zalaczniki": "Załączniki",
    "zalaczniki": "załączniki",
    "plikow": "plików",
    "materialy": "materiały",
    "materialow": "materiałów",
    "przeplywu": "przepływu",
    "watku": "wątku",
    "pomoca": "pomocą",
    "mechanizmow": "mechanizmów",
    "schematow": "schematów",
    "umozliwiajacych": "umożliwiających",
    "przejscia": "przejścia",
    "zostal": "został",
    "zostala": "została",
    "zostaly": "zostały",
    "polaczyc": "połączyć",
    "niewielka": "niewielką",
    "plytke": "płytkę",
    "plytki": "płytki",
    "sterujaca": "sterującą",
    "sztucznego": "sztucznego",
    "horyzontu": "horyzontu",
    "czytelnej": "czytelnej",
    "fizycznego": "fizycznego",
    "przenosnym": "przenośnym",
    "układem": "układem",
    "biblioteke": "bibliotekę",
    "autorska": "autorską",
    "pozwalajaca": "pozwalającą",
    "dzialajacy": "działający",
    "dzialajaca": "działająca",
    "dzialanie": "działanie",
    "aktywny": "aktywny",
    "Sciezka": "Ścieżka",
    "sciezka": "ścieżka",
    "wywolan": "wywołań",
    "wywolanie": "wywołanie",
    "wywoluje": "wywołuje",
    "wywolac": "wywołać",
    "przygotowany": "przygotowany",
    "przygotowana": "przygotowana",
    "kompensacji": "kompensacji",
    "kierunek": "kierunek",
    "wektora": "wektora",
    "nieba": "nieba",
    "ziemi": "ziemi",
    "wielokaty": "wielokąty",
    "dluzsze": "dłuższe",
    "krotsze": "krótsze",
    "niepotrzebnego": "niepotrzebnego",
    "obciazenia": "obciążenia",
    "zweryfikowac": "zweryfikować",
    "potwierdzic": "potwierdzić",
    "uzupelnic": "uzupełnić",
    "stabilnosc": "stabilność",
    "stabilne": "stabilne",
    "sprawdzic": "sprawdzić",
    "sprawdzonych": "sprawdzonych",
    "dopracowaniu": "dopracowaniu",
    "Urzadzenie": "Urządzenie",
    "urzadzenie": "urządzenie",
    "zastapic": "zastąpić",
    "odkomentowaniu": "odkomentowaniu",
    "elementow": "elementów",
    "niezaleznym": "niezależnym",
    "zlozony": "złożony",
    "zlozona": "złożona",
    "opoznienie": "opóźnienie",
    "opoznien": "opóźnień",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C4D1", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Strona ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.style = "Footer"
    footer.add_run("Dokumentacja projektu AHRS")
    footer.add_run(" " * 58)
    add_page_number(footer)


def add_caption(doc, label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(label)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)


def add_figure(doc, path, caption, width_cm=14.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.style = "CodeBlock"
    p.add_run(text)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def polish_text(text):
    for src, dst in sorted(POLISH_REPLACEMENTS.items(), key=lambda item: len(item[0]), reverse=True):
        text = text.replace(src, dst)
    for src, dst in sorted(POST_POLISH_REPLACEMENTS.items(), key=lambda item: len(item[0]), reverse=True):
        text = text.replace(src, dst)
    return text


def apply_polish_diacritics(doc):
    # Code snippets and literal paths remain untouched; prose, headings, captions and tables get Polish glyphs.
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "CodeBlock":
            continue
        for run in paragraph.runs:
            run.text = polish_text(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.style.name == "CodeBlock":
                        continue
                    for run in paragraph.runs:
                        run.text = polish_text(run.text)


def add_kv_table(doc, rows, widths=(2700, 6660)):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_width(cells[0], widths[0])
        set_cell_width(cells[1], widths[1])
        set_cell_shading(cells[0], "E8EEF5")
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[0].paragraphs[0].add_run(key).bold = True
        cells[1].paragraphs[0].add_run(value)
    return table


def add_matrix_table(doc, title, matrix_rows):
    doc.add_paragraph(title, style="CaptionBeforeTable")
    table = doc.add_table(rows=1, cols=len(matrix_rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="AEB8C2", size="4")
    for i, value in enumerate(matrix_rows[0]):
        table.rows[0].cells[i].text = value
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row in matrix_rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8.5)
    return table


def draw_architecture_diagram():
    out = IMG / "AHRS_architektura_systemu.png"
    w, h = 1800, 980
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arial.ttf", 46)
        box_font = ImageFont.truetype("arial.ttf", 30)
        small_font = ImageFont.truetype("arial.ttf", 24)
    except OSError:
        title_font = box_font = small_font = ImageFont.load_default()

    palette = {
        "blue": "#E6F0FA",
        "blue_border": "#5D7FA3",
        "green": "#EAF5EA",
        "green_border": "#6C976A",
        "orange": "#FFF1DE",
        "orange_border": "#B78643",
        "gray": "#F3F5F7",
        "gray_border": "#8A98A8",
    }

    def box(x, y, bw, bh, title, lines, fill, border):
        draw.rounded_rectangle((x, y, x + bw, y + bh), radius=22, fill=fill, outline=border, width=4)
        draw.text((x + 26, y + 22), title, fill="#1C2733", font=box_font)
        yy = y + 72
        for line in lines:
            draw.text((x + 26, yy), line, fill="#334155", font=small_font)
            yy += 34

    def arrow(x1, y1, x2, y2, text=None):
        draw.line((x1, y1, x2, y2), fill="#334155", width=5)
        angle = 18
        if x2 >= x1:
            pts = [(x2, y2), (x2 - 28, y2 - angle), (x2 - 28, y2 + angle)]
        else:
            pts = [(x2, y2), (x2 + 28, y2 - angle), (x2 + 28, y2 + angle)]
        draw.polygon(pts, fill="#334155")
        if text:
            draw.text(((x1 + x2) / 2 - 80, (y1 + y2) / 2 - 42), text, fill="#475569", font=small_font)

    draw.text((70, 50), "Architektura przeplywu danych w projekcie AHRS", fill="#0B2545", font=title_font)
    box(70, 170, 360, 210, "Czujniki IMU", ["LSM6DS3TR-C", "akcelerometr + zyroskop", "LIS3MDL magnetometr"], palette["green"], palette["green_border"])
    box(520, 170, 380, 210, "Warstwa IMU", ["kalibracja EEPROM", "mapowanie osi", "pomiary ax, ay, az", "gx, gy, gz, mx, my, mz"], palette["blue"], palette["blue_border"])
    box(990, 170, 380, 210, "Fuzja orientacji", ["Adafruit NXP Fusion", "lub wlasny filtr Kalmana", "roll, pitch, heading"], palette["orange"], palette["orange_border"])
    box(1430, 170, 300, 210, "Watek AHRS", ["30 Hz", "DataPoint", "mutex FreeRTOS"], palette["gray"], palette["gray_border"])
    box(420, 570, 460, 210, "Watek wyswietlacza", ["prog zmian: 0,2 deg / 1 deg", "odswiezanie 40 Hz", "bufory GFXcanvas16"], palette["blue"], palette["blue_border"])
    box(1010, 570, 440, 210, "Interfejs uzytkownika", ["sztuczny horyzont", "drabinka pitch", "skala przechylenia", "kompas polkolisty"], palette["orange"], palette["orange_border"])

    arrow(430, 275, 520, 275, "I2C")
    arrow(900, 275, 990, 275)
    arrow(1370, 275, 1430, 275)
    arrow(1580, 380, 1250, 570, "roll/pitch/heading")
    arrow(880, 675, 1010, 675)
    arrow(650, 570, 650, 410, "wspolne dane")

    img.save(out, quality=95)
    return out


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(12)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = RGBColor(76, 91, 109)
    subtitle.paragraph_format.space_after = Pt(18)

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "CodeBlock" not in styles:
        code = styles.add_style("CodeBlock", 1)
    else:
        code = styles["CodeBlock"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor(30, 41, 59)
    code.paragraph_format.left_indent = Cm(0.3)
    code.paragraph_format.right_indent = Cm(0.3)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0

    if "CaptionBeforeTable" not in styles:
        cap = styles.add_style("CaptionBeforeTable", 1)
    else:
        cap = styles["CaptionBeforeTable"]
    cap.font.name = "Calibri"
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor(80, 80, 80)
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(4)


def add_cover(doc):
    p = doc.add_paragraph("Krakow, 18.05.2026")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(70)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Dokumentacja projektu AHRS")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("System wyznaczania orientacji przestrzennej z wyswietlaczem sztucznego horyzontu")

    doc.add_paragraph()
    meta = add_kv_table(doc, [
        ("Projekt", "AHRS - Attitude and Heading Reference System"),
        ("Platforma", "Arduino Nano ESP32 / FreeRTOS / C++"),
        ("Aktywna fuzja danych", "Biblioteka Adafruit AHRS, wariant Adafruit_NXPSensorFusion"),
        ("Wariant alternatywny", "Autorski filtr Kalmana przygotowany w src/AHRS/src/kalmanFilter"),
        ("Podstawa opracowania", "Kod z repozytorium, schematy elektryczne, rysunki techniczne oraz obrazy w katalogu projektu"),
    ])
    for row in meta.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Dokument przygotowany w ukladzie sprawozdania technicznego, wzorowanym na przykladowej dokumentacji projektu.")
    run.italic = True
    run.font.color.rgb = RGBColor(80, 80, 80)
    doc.add_page_break()


def add_static_toc(doc):
    doc.add_heading("Spis tresci", level=1)
    items = [
        "1. Wprowadzenie i cele",
        "2. Wykorzystane komponenty i urzadzenia",
        "3. Konstrukcja mechaniczna i schemat elektryczny",
        "4. Oprogramowanie i architektura projektu",
        "5. Akwizycja danych i aktualna fuzja Adafruit AHRS",
        "6. Autorski filtr Kalmana",
        "7. Interfejs uzytkownika na wyswietlaczu",
        "8. Uruchomienie, konfiguracja i testowanie",
        "9. Podsumowanie i dalsze kierunki rozwoju",
        "10. Zalaczniki: struktura plikow i wykorzystane materialy",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.2)
        p.add_run(item)
    doc.add_page_break()


def add_main_content(doc):
    arch = draw_architecture_diagram()

    doc.add_heading("1. Wprowadzenie i cele", level=1)
    doc.add_paragraph(
        "Projekt AHRS (Attitude and Heading Reference System) jest przenosnym ukladem do "
        "wyznaczania orientacji przestrzennej oraz prezentowania jej w postaci sztucznego "
        "horyzontu i kompasu. Urzadzenie zbiera dane z akcelerometru, zyroskopu i "
        "magnetometru, wykonuje fuzje sensoryczna, a wynikowy kat pochylenia, przechylenia "
        "oraz kurs wyswietla na ekranie TFT."
    )
    doc.add_paragraph(
        "W aktualnej wersji programu aktywna jest biblioteka Adafruit AHRS, a dokladniej "
        "klasa Adafruit_NXPSensorFusion. Repozytorium zawiera rowniez autorska biblioteke "
        "filtru Kalmana, ktora moze zastapic gotowy filtr po odkomentowaniu wywolan "
        "kalmanSetup(...) oraz kalmanRun(...). Z tego powodu dokumentacja opisuje oba "
        "warianty: dzialajacy tor Adafruit oraz przygotowany tor Kalmana wraz z rownaniami."
    )
    add_figure(doc, arch, "Rys. 1 Uproszczona architektura przeplywu danych w systemie AHRS", 15.5)
    doc.add_paragraph("Glowne cele projektu:")
    add_bullets(doc, [
        "odczyt orientacji przestrzennej z zestawu IMU 9-DOF;",
        "prezentacja roll, pitch i heading w formie czytelnej dla operatora;",
        "oddzielenie watku fuzji danych od watku wyswietlania za pomoca mechanizmow FreeRTOS;",
        "zachowanie mozliwosci przejscia z gotowej biblioteki Adafruit AHRS na autorski filtr Kalmana;",
        "opracowanie obudowy oraz schematow elektrycznych umozliwiajacych wykonanie fizycznego prototypu.",
    ])

    doc.add_heading("2. Wykorzystane komponenty i urzadzenia", level=1)
    doc.add_paragraph(
        "Zawartosc katalogu Docs/img wskazuje zestaw elementow przewidzianych do budowy "
        "urzadzenia. Zestaw zostal dobrany tak, aby polaczyc niewielka plytke sterujaca, "
        "czujnik inercyjny i kolorowy wyswietlacz z niezaleznym ukladem zasilania."
    )
    components = [
        ("Arduino Nano ESP32", "Glowny mikrokontroler uruchamiajacy szkic AHRS.ino i zadania FreeRTOS."),
        ("IMU 9-DOF", "Zestaw LSM6DS3TR-C oraz LIS3MDL: akcelerometr, zyroskop i magnetometr."),
        ("Wyswietlacz ST7789 240 x 320", "Ekran do prezentacji sztucznego horyzontu i polkolistego kompasu."),
        ("Adapter zaciskowy Nano", "Ulatwia wyprowadzenie pinow Arduino do prototypu i obudowy."),
        ("Przetwornica step-up", "Element toru zasilania, widoczny rowniez w schemacie PowerSupply."),
        ("Przelacznik suwakowy", "Mechaniczny element wlaczania lub odcinania zasilania urzadzenia."),
        ("Obudowa drukowana 3D", "Modele STL i SolidWorks w katalogu Case/3dPrint oraz Case/3dModels."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    headers = table.rows[0].cells
    headers[0].text = "Komponent"
    headers[1].text = "Rola w projekcie"
    set_repeat_table_header(table.rows[0])
    for cell in headers:
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for name, desc in components:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = desc
    for row in table.rows:
        set_cell_width(row.cells[0], 2700)
        set_cell_width(row.cells[1], 6660)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_caption(doc, "Tab. 1 Zestawienie podstawowych elementow systemu AHRS")
    add_figure(doc, IMG / "Arduino_Nano_ESP32.png", "Rys. 2 Plytka Arduino Nano ESP32 wykorzystana jako jednostka sterujaca", 9.0)
    add_figure(doc, IMG / "IMU.png", "Rys. 3 Modul IMU dostarczajacy pomiary inercyjne i magnetyczne", 7.0)
    add_figure(doc, IMG / "LCD.png", "Rys. 4 Wyswietlacz TFT uzywany do prezentacji wskazan AHRS", 7.0)

    doc.add_heading("3. Konstrukcja mechaniczna i schemat elektryczny", level=1)
    doc.add_paragraph(
        "Warstwa sprzetowa projektu zostala rozdzielona na schemat plytki Arduino, tor "
        "zasilania oraz obszar prototypowy. Taki podzial ulatwia weryfikacje polaczen: "
        "osobno mozna analizowac linie mikrokontrolera, przetwornice i elementy wykonawcze "
        "montowane na plytce prototypowej."
    )
    add_figure(doc, SCHEMATICS / "ArduinoBoard.png", "Rys. 5 Schemat polaczen plytki Arduino Nano ESP32", 15.0)
    add_figure(doc, SCHEMATICS / "PowerSupply.png", "Rys. 6 Schemat toru zasilania ukladu", 11.0)
    add_figure(doc, SCHEMATICS / "Prototyping_Area.png", "Rys. 7 Schemat obszaru prototypowego", 15.0)
    doc.add_paragraph(
        "W katalogu Case znajduja sie modele SolidWorks oraz pliki STL przygotowane do "
        "druku 3D: Case.STL, Front.STL oraz BatteryCover.STL. Rysunek techniczny obudowy "
        "w pliku Obudowa.pdf wskazuje miedzy innymi gabaryty 140 mm x 75 mm, przekroje A-A "
        "i B-B oraz skale 1:2. Dokumentacja mechaniczna pelni role wykonawcza: pozwala "
        "odtworzyc obudowe, front oraz pokrywe baterii."
    )
    add_kv_table(doc, [
        ("Modele 3D", "Case/3dModels: AHRS.SLDASM, Case.SLDPRT, Front.SLDPRT, BatteryCover.SLDPRT"),
        ("Pliki do druku", "Case/3dPrint: Case.STL, Front.STL, BatteryCover.STL"),
        ("Rysunek wykonawczy", "Case/Technical_Drawings/Obudowa.pdf oraz AHRS.SLDDRW"),
    ])

    doc.add_heading("4. Oprogramowanie i architektura projektu", level=1)
    doc.add_paragraph(
        "Kod programu znajduje sie w katalogu src/AHRS. Glowny plik AHRS.ino tworzy obiekty "
        "Display, IMU, DeviceControl oraz KalmanFilter, a nastepnie uruchamia dwa zadania "
        "FreeRTOS przypiete do roznych rdzeni mikrokontrolera."
    )
    add_code(doc, """xTaskCreatePinnedToCore(taskAHRS, "AHRS", 8192, NULL, 2, NULL, 1);
xTaskCreatePinnedToCore(taskDisplay, "DISPLAY", 16384, NULL, 1, NULL, 0);""")
    doc.add_paragraph(
        "Zadanie AHRS odpowiada za odczyt czujnikow i wyznaczenie orientacji. Zadanie DISPLAY "
        "pracuje niezaleznie, pobiera ostatnie dane przez mutex i odswieza tylko te elementy "
        "ekranu, ktore zmienily sie ponad prog."
    )
    add_kv_table(doc, [
        ("taskAHRS", "Odczyt IMU, fuzja danych, zapis DataPoint do wspolnej struktury przez attitudeMutex."),
        ("taskDisplay", "Porownanie zmian katow, aktualizacja sztucznego horyzontu i kompasu."),
        ("IMU", "Inicjalizacja LSM6DS3TR-C i LIS3MDL, kalibracja, transformacja osi, wywolanie filtru."),
        ("Display", "Rysowanie horyzontu, drabinki pitch, skali przechylenia, sylwetki samolotu i kompasu."),
        ("KalmanFilter", "Predykcja i aktualizacja stanu z uzyciem klas Matrix oraz Vector."),
        ("linAlg", "Minimalna algebra liniowa: mnozenie macierzy, transpozycja, wyznacznik i macierz odwrotna."),
    ])
    doc.add_paragraph("Najwazniejsze stale konfiguracyjne:")
    add_bullets(doc, [
        "FILTER_UPDATE_RATE_HZ = 30 - czestotliwosc pracy filtru orientacji;",
        "ATT_REFRESH_RATE_HZ = 40 - odswiezanie warstwy wyswietlania;",
        "ATT_THRESHOLD_DEG = 0.2 - minimalna zmiana roll/pitch wymuszajaca przerysowanie horyzontu;",
        "HDG_THRESHOLD_DEG = 1.0 - minimalna zmiana kursu wymuszajaca przerysowanie kompasu;",
        "PIXELS_PER_DEGREE = 7.0 - przelicznik kata pitch na przesuniecie horyzontu.",
    ])

    doc.add_heading("5. Akwizycja danych i aktualna fuzja Adafruit AHRS", level=1)
    doc.add_paragraph(
        "Klasa IMU inicjalizuje dwa uklady Adafruit: LSM6DS3TR-C oraz LIS3MDL. Po starcie "
        "ladowana jest kalibracja z EEPROM, ustawiane sa zakresy pomiarowe i czestotliwosci "
        "probkowania. Akcelerometr i zyroskop pracuja z czestotliwoscia 104 Hz, magnetometr "
        "w trybie ciaglym z szybkim odczytem, a magistrala I2C zostaje ustawiona na 400 kHz."
    )
    add_code(doc, """lsm6ds.setAccelRange(LSM6DS_ACCEL_RANGE_2_G);
lsm6ds.setGyroRange(LSM6DS_GYRO_RANGE_250_DPS);
lis3mdl.setRange(LIS3MDL_RANGE_4_GAUSS);
filter.begin(FILTER_UPDATE_RATE_HZ);
Wire.setClock(400000);""")
    doc.add_paragraph(
        "W funkcji IMU::read(...) dane sa kalibrowane, a nastepnie osie sa przemapowane do "
        "ukladu przyjetego w projekcie. Aktywna sciezka wywoluje filter.update(...) z biblioteki "
        "Adafruit i pobiera roll, pitch oraz yaw. Heading jest dodatkowo przesuniety o 90 stopni "
        "i zawiniety operacja fmod(..., 360)."
    )
    add_code(doc, """filter.update(gx, gy, gz, ax, ay, az, mx, my, mz);
dataPoint.roll = filter.getRoll();
dataPoint.pitch = filter.getPitch();
dataPoint.heading = fmod(-filter.getYaw() - 90.0f, 360.0f);""")

    doc.add_heading("6. Autorski filtr Kalmana", level=1)
    doc.add_paragraph(
        "Autorska implementacja filtru znajduje sie w katalogu src/AHRS/src/kalmanFilter. "
        "Obecnie jest ona przygotowana jako alternatywa dla biblioteki Adafruit: obiekt "
        "KalmanFilter jest tworzony w AHRS.ino, ale wywolanie kalmanRun(...) w zadaniu AHRS "
        "oraz kalmanSetup(...) w DeviceControl::powerOn(...) pozostaja zakomentowane."
    )
    add_code(doc, """// localDataPoint = kalmanRun(kalmanFilter, measurement);
// kalmanSetup(kalmanData);""")
    doc.add_heading("6.1. Wektor stanu, sterowanie i pomiar", level=2)
    doc.add_paragraph(
        "Stan filtru jest szesciowymiarowy. Pierwsze trzy skladowe reprezentuja katy Eulera, "
        "a kolejne trzy skladowe reprezentuja estymowane skladowe dryfu zyroskopu:"
    )
    add_code(doc, """x_hat = [ roll, pitch, yaw, b_gx, b_gy, b_gz ]^T
u     = [ gx, gy, gz ]^T
z     = [ roll_acc, pitch_acc, yaw_mag ]^T""")
    doc.add_paragraph(
        "Pomiary pomocnicze sa wyznaczane z akcelerometru i magnetometru. Dla roll i pitch "
        "uzywany jest kierunek wektora grawitacji, natomiast yaw powstaje po kompensacji "
        "pochylenia magnetometru:"
    )
    add_code(doc, """mxh = mx*cos(theta) + mz*sin(theta)
myh = mx*sin(phi)*sin(theta) + my*cos(phi) - mz*sin(phi)*cos(theta)

roll_acc  = atan2(ay, az)
pitch_acc = atan2(-ax, sqrt(ay^2 + az^2))
yaw_mag   = atan2(-myh, mxh)""")

    doc.add_heading("6.2. Macierze modelu", level=2)
    doc.add_paragraph(
        "Dla czestotliwosci FILTER_UPDATE_RATE_HZ = 30 przyjeto dt = 1/30 s. Macierz F "
        "utrzymuje poprzedni stan i modeluje wplyw dryfu zyroskopu na katy, a macierz G "
        "wprowadza pomiar predkosci katowej do pierwszych trzech skladowych stanu."
    )
    add_matrix_table(doc, "Tab. 2 Macierz przejscia F uzyta w kalmanSetup(...)", [
        ["1", "0", "0", "-dt", "0", "0"],
        ["0", "1", "0", "0", "-dt", "0"],
        ["0", "0", "1", "0", "0", "-dt"],
        ["0", "0", "0", "1", "0", "0"],
        ["0", "0", "0", "0", "1", "0"],
        ["0", "0", "0", "0", "0", "1"],
    ])
    add_matrix_table(doc, "Tab. 3 Macierz sterowania G", [
        ["dt", "0", "0"],
        ["0", "dt", "0"],
        ["0", "0", "dt"],
        ["0", "0", "0"],
        ["0", "0", "0"],
        ["0", "0", "0"],
    ])
    add_matrix_table(doc, "Tab. 4 Macierz obserwacji H", [
        ["1", "0", "0", "0", "0", "0"],
        ["0", "1", "0", "0", "0", "0"],
        ["0", "0", "1", "0", "0", "0"],
    ])
    doc.add_paragraph(
        "Macierz R opisuje niepewnosc pomiarow roll, pitch i yaw: sigma_r = 0,15, "
        "sigma_p = 0,15 oraz sigma_y = 3,0. Macierz Q rozdziela szum procesu na skladowe "
        "katow qAngle = 0,001 oraz dryfu qBias = 0,0003."
    )

    doc.add_heading("6.3. Rownania predykcji i aktualizacji", level=2)
    doc.add_paragraph(
        "Implementacja KalmanFilter.cpp odpowiada klasycznemu dyskretnemu filtrowi Kalmana. "
        "Predykcja sklada sie z przewidywania stanu oraz kowariancji:"
    )
    add_code(doc, """x_hat_k^- = F * x_hat_(k-1) + G * u_k
P_k^-     = F * P_(k-1) * F^T + Q""")
    doc.add_paragraph(
        "Aktualizacja wykorzystuje innowacje pomiarowa z_k - H*x_hat_k^- oraz wzmocnienie K:"
    )
    add_code(doc, """K_k       = P_k^- * H^T * inv(H * P_k^- * H^T + R)
x_hat_k   = x_hat_k^- + K_k * (z_k - H * x_hat_k^-)
P_k       = (I - K_k*H) * P_k^- * (I - K_k*H)^T + K_k * R * K_k^T""")
    doc.add_paragraph(
        "Ostatnie rownanie jest postacia Josepha aktualizacji kowariancji. Jest bardziej "
        "odporne numerycznie niz proste P = (I - K*H)P, poniewaz jawnie uwzglednia szum "
        "pomiaru i pomaga zachowac symetrie oraz dodatnia polokreslonosc macierzy P."
    )
    doc.add_paragraph(
        "Przy ponownym wlaczaniu autorskiego filtru nalezy zadbac o spojnosc jednostek. "
        "W aktualnej sciezce Adafruit zyroskop jest przekazywany w stopniach na sekunde, "
        "natomiast roll_acc, pitch_acc i yaw_mag w kalmanRun(...) sa liczone w radianach. "
        "Najbezpieczniejszym wariantem jest przeliczenie gx, gy, gz na rad/s albo "
        "konsekwentne prowadzenie calego filtru w stopniach."
    )

    doc.add_heading("7. Interfejs uzytkownika na wyswietlaczu", level=1)
    doc.add_paragraph(
        "Warstwa Display rysuje caly interfejs w pamieci za pomoca buforow GFXcanvas16, a "
        "nastepnie przesyla gotowa bitmape do ekranu. Gorna czesc wyswietlacza 240 x 240 "
        "px zajmuje sztuczny horyzont, a dolna czesc 240 x 80 px tworzy kompas."
    )
    add_figure(doc, IMG / "Attitude_Indicator_Level.png", "Rys. 8 Widok sztucznego horyzontu w locie poziomym", 8.5)
    add_figure(doc, IMG / "Attitude_Indicatur_Banked.png", "Rys. 9 Widok sztucznego horyzontu przy przechyleniu", 8.5)
    doc.add_paragraph(
        "Funkcja showAttitude(...) przelicza roll na wektory kierunkowe, przesuwa horyzont "
        "wedlug pitch, rysuje tlo nieba i ziemi, linie drabinki pitch co 2,5 stopnia, skale "
        "przechylenia oraz staly symbol samolotu. Funkcja showCompass(...) rysuje gorna "
        "polowe podzialki kursu i etykiety N, E, S, W."
    )
    add_bullets(doc, [
        "niebo i ziemia sa rysowane jako wypelnione wielokaty, co daje plynna zmiane przechylenia;",
        "dluzsze znaczniki pitch pojawiaja sie co 10 stopni, krotsze co 5 i 2,5 stopnia;",
        "aktualizacja ekranu jest ograniczana progami, aby uniknac niepotrzebnego migotania i obciazenia SPI.",
    ])

    doc.add_heading("8. Uruchomienie, konfiguracja i testowanie", level=1)
    doc.add_paragraph(
        "Po zmontowaniu ukladu procedura uruchomienia powinna obejmowac osobna weryfikacje "
        "zasilania, komunikacji I2C, inicjalizacji czujnikow oraz poprawnosci wskazan na ekranie."
    )
    add_numbered(doc, [
        "Sprawdzic tor zasilania zgodnie ze schematem PowerSupply oraz poprawna pozycje przelacznika.",
        "Podlaczyc Arduino Nano ESP32 i zweryfikowac, czy linie LCD_CS, TP_CS i LCD_BCKLIT sa ustawione zgodnie z display.hpp.",
        "Uruchomic monitor portu szeregowego 9600 baud i potwierdzic komunikaty inicjalizacji LSM6DS oraz LIS3MDL.",
        "Wykonac kalibracje czujnikow Adafruit i zapisac ja w EEPROM, aby cal.loadCalibration() moglo ja odczytac przy starcie.",
        "Porownac zmiane roll, pitch i heading z fizycznym ruchem obudowy.",
        "Jesli wlaczany jest filtr Kalmana, odkomentowac kalmanSetup(...) i kalmanRun(...), po czym sprawdzic jednostki oraz stabilnosc macierzy P.",
    ])
    doc.add_paragraph(
        "Weryfikacja programowa powinna obejmowac przynajmniej testy algebry liniowej, poniewaz "
        "od poprawnego mnozenia, transpozycji i odwracania macierzy zalezy stabilnosc filtru. "
        "Plik CMakeLists.txt odwoluje sie do test/kalmanFilter, ale taki katalog nie znajduje "
        "sie w aktualnej zawartosci repozytorium, dlatego warto go uzupelnic przed dalszym rozwojem."
    )

    doc.add_heading("9. Podsumowanie i dalsze kierunki rozwoju", level=1)
    doc.add_paragraph(
        "Projekt AHRS laczy kompletna warstwe sprzetowa, obudowe, schematy elektryczne oraz "
        "oprogramowanie realizujace fuzje danych i wizualizacje orientacji. Aktualna wersja "
        "korzysta ze sprawdzonej biblioteki Adafruit AHRS, co pozwala szybko uzyskac stabilne "
        "wskazania na prototypie. Rownolegle w repozytorium istnieje autorski filtr Kalmana, "
        "ktory moze stac sie docelowym algorytmem po dopracowaniu jednostek, inicjalizacji i testow."
    )
    doc.add_paragraph("Proponowane dalsze prace:")
    add_bullets(doc, [
        "uruchomienie autorskiego filtru Kalmana w trybie porownawczym obok Adafruit AHRS;",
        "dodanie testow jednostkowych Matrix, Vector oraz rownan filtru;",
        "zapis danych orientacji do portu szeregowego lub karty pamieci w celu analizy dryfu;",
        "dodanie procedury kalibracji magnetometru dostepnej z poziomu interfejsu uzytkownika;",
        "przygotowanie finalnej plytki PCB na podstawie istniejacych schematow prototypowych.",
    ])

    doc.add_heading("10. Zalaczniki: struktura plikow i wykorzystane materialy", level=1)
    add_kv_table(doc, [
        ("Kod glowny", "src/AHRS/AHRS.ino"),
        ("Czujniki", "src/AHRS/src/sensor/imu.cpp, imu.hpp"),
        ("Wyswietlacz", "src/AHRS/src/display/display.cpp, display.hpp"),
        ("Sterowanie urzadzeniem", "src/AHRS/src/deviceControl/deviceControl.cpp, deviceControl.hpp"),
        ("Filtr Kalmana", "src/AHRS/src/kalmanFilter/KalmanFilter.cpp, KalmanFilter.hpp"),
        ("Algebra liniowa", "src/AHRS/src/linAlg/Matrix.cpp, Vector.cpp oraz pliki naglowkowe"),
        ("Schematy elektryczne", "Electrical_schematics/ArduinoBoard.*, PowerSupply.*, Prototyping_Area.*"),
        ("Grafiki komponentow", "Docs/img/*.png oraz Docs/img/*.jpg"),
        ("Obudowa", "Case/3dModels, Case/3dPrint, Case/Technical_Drawings"),
    ])


def build():
    doc = Document()
    configure_styles(doc)
    add_footer(doc.sections[0])
    add_cover(doc)
    add_static_toc(doc)
    add_main_content(doc)
    apply_polish_diacritics(doc)
    doc.core_properties.title = "Dokumentacja projektu AHRS"
    doc.core_properties.subject = "AHRS, Adafruit AHRS, filtr Kalmana"
    doc.core_properties.comments = "Wygenerowano na podstawie kodu i artefaktow projektu AHRS."
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
